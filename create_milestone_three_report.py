from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

report_path = "Milestone_Three_Exemplary_Report.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Milestone Three Report")
title_run.bold = True
title_run.font.size = Pt(18)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("DX699 - AI for Leaders\nBivariate and Time-Series Analysis of Social Media Advertising Performance\nAwais Elahi")

doc.add_paragraph()

sections = [
    ("Executive Summary", [
        "This report applies the analysis and visualization techniques covered through Weeks 5 to 7 to the Social Media Advertising dataset. The goal is to identify which campaign attributes are most associated with return on investment, evaluate whether the data are usable for decision-making, and present the results in a way that supports managerial action.",
        "The analysis shows that ROI is more strongly associated with engagement behavior than with raw volume measures. Engagement Score has the strongest positive linear relationship with ROI among the variables tested, while Clicks and Impressions show weaker positive associations. Conversion Rate and Acquisition Cost show almost no linear relationship with ROI in this sample, suggesting that simple pairwise relationships may not explain campaign performance on their own.",
        "The time-based analysis of monthly average ROI suggests that performance is relatively stable over the observed period, with only modest month-to-month movement. This indicates that the dataset is usable for exploratory analysis, but it also suggests that more detailed segmentation by channel, audience, and campaign goal may be needed to uncover stronger drivers of performance."
    ]),
    ("1. Business Problem and Objective", [
        "The central question for this milestone is which observable campaign characteristics are most useful for understanding advertising performance. For a decision-maker, the practical issue is not just whether campaigns attract attention, but which features are most closely connected to financial outcomes such as ROI.",
        "The objective of this report is to evaluate pairwise relationships between key campaign variables, identify whether some variables appear redundant, examine whether campaign performance changes over time, and demonstrate clear, decision-oriented visual communication."
    ]),
    ("2. Dataset Description", [
        "The report uses the Social Media Advertising dataset, which includes campaign-level observations such as Conversion Rate, Acquisition Cost, ROI, Clicks, Impressions, Engagement Score, Channel Used, Campaign Goal, Customer Segment, and Date.",
        "The data appear suitable for exploratory analysis because they contain a combination of performance, cost, audience, and timing variables. This supports both cross-sectional analysis and simple time-based trend analysis."
    ]),
    ("3. Methods Applied", [
        "Bivariate analysis was used to evaluate the relationship between ROI and other numerical variables. Correlations were calculated to understand the strength and direction of linear relationships.",
        "A pair plot was produced for Conversion Rate, ROI, Clicks, and Engagement Score. This provided a compact way to compare distributions and pairwise patterns across several variables at once.",
        "Monthly average ROI was computed and plotted to evaluate whether advertising performance drifted over time.",
        "Visual emphasis techniques from the Storytelling With Data exercises were applied to improve clarity, reduce clutter, and direct viewer attention toward the most decision-relevant message."
    ]),
    ("4. Key Findings", [
        "The strongest pairwise relationship with ROI among the numerical variables examined was Engagement Score.",
        "The pair plot indicates that Clicks and Impressions move together more closely than either one moves with ROI. That pattern suggests partial redundancy.",
        "The monthly average ROI line remains within a narrow band across the observed year, which implies there is no dramatic time drift in the current sample.",
        "Week 7 work showed how preattentive attributes such as color, contrast, shading, and line thickness can direct attention to the most important segment of a chart."
    ]),
    ("5. Interpretation and Business Implications", [
        "These results suggest that stronger engagement is a more promising signal of financial performance than simple campaign scale measures. A manager reviewing new campaign proposals should therefore pay attention to whether a campaign is likely to generate meaningful audience interaction, not just high reach.",
        "The lack of a strong linear relationship between Conversion Rate and ROI in this sample does not necessarily mean conversion is unimportant. It may indicate that ROI is influenced by multiple interacting factors, such as campaign goal, audience type, platform, or content quality.",
        "The stable monthly ROI pattern is encouraging from a reporting standpoint. It suggests the dataset is usable for near-term benchmarking, but it also implies that richer segmentation may be more informative than simple overall averages."
    ]),
    ("6. Limitations", [
        "This analysis focuses mostly on pairwise relationships, which can miss nonlinear or interaction effects.",
        "Correlation does not imply causation, so the findings should not be treated as proof of what drives ROI.",
        "Important confounding variables may exist, including creative quality, channel strategy, brand strength, and audience targeting.",
        "The monthly trend is based on average ROI and may conceal important variation within specific channels or customer segments."
    ]),
    ("7. Recommendations", [
        "Segment the data by channel and customer segment to test whether relationships differ across groups.",
        "Compare ROI across campaign goals and platforms rather than relying only on overall averages.",
        "Build multivariate models in later milestones to test whether combinations of engagement, clicks, and audience variables explain ROI better than pairwise analysis alone.",
        "Continue using direct-label, highlighted, low-clutter visualizations so that decision-makers can identify the central finding quickly."
    ]),
    ("8. Conclusion", [
        "The Weeks 5 to 7 work provides a useful foundation for business-focused data analysis. In this milestone, bivariate analysis, pair plots, time-series plots, and narrative visualization techniques were applied to a real campaign dataset.",
        "The results show that engagement-based measures appear more informative for ROI than raw exposure measures alone, that some metrics may be partially redundant, and that average ROI is relatively stable over time.",
        "Overall, the dataset is suitable for continued analysis, and the findings provide a reasonable basis for moving toward deeper multivariate and decision-support modeling in the next phase of the project."
    ]),
]

for heading, paragraphs in sections:
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)


doc.add_heading("Correlation Summary", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Light List Accent 1"
header = table.rows[0].cells
header[0].text = "Variable"
header[1].text = "Correlation with ROI"
for variable, value in [
    ("ROI", "1.000"),
    ("Engagement Score", "0.355"),
    ("Clicks", "0.188"),
    ("Impressions", "0.166"),
    ("Conversion Rate", "-0.000"),
    ("Acquisition Cost", "-0.002"),
]:
    row = table.add_row().cells
    row[0].text = variable
    row[1].text = value

doc.add_heading("Appendix: Weekly Work Completed", level=1)
doc.add_paragraph("Week 5: Scatter plots, pair plots, correlations, line plots, area plots, and storytelling-style charts were completed. Missing reference images were replaced with generated plots so the homework section runs end to end.")
doc.add_paragraph("Week 6: A written response comparing area plots with separate line charts was completed. A dataset analysis using the Social Media Advertising dataset was added, along with a storytelling-style funnel drop-off chart.")
doc.add_paragraph("Week 7: A written response on using preattentive attributes to emphasize the last three months of a time series was completed. A highlighted time-series example and a storytelling-style before/after chart were created.")

doc.save(report_path)
print(f"Created {report_path}")
