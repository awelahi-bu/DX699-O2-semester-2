import base64
from pathlib import Path

import nbformat
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "Milestone_Three_Final_Polished_Report.docx"
REPORT_OUT = ROOT / "Milestone_Three_Final_Polished_Report_With_Graphs.docx"
FIG_DIR = ROOT / "report_figures"
FIG_DIR.mkdir(exist_ok=True)

# Notebook targets and the code signature to locate relevant plotted cells.
FIG_SPECS = [
    {
        "notebook": ROOT / "Week5-concealed-answer.ipynb",
        "match": "sns.pairplot(df)",
        "filename": "week5_pairplot.png",
        "caption": "Figure 1. Week 5 pair plot showing key bivariate relationships."
    },
    {
        "notebook": ROOT / "OMDS-MODB2-Week6-Elahi-Awais.ipynb",
        "match": "df.plot.area()",
        "filename": "week6_area_plot.png",
        "caption": "Figure 2. Week 6 area plot comparing nitrate and phosphate trends."
    },
    {
        "notebook": ROOT / "OMDS-MODB2-Week6-Elahi-Awais.ipynb",
        "match": "Average ROI by Month",
        "filename": "week6_roi_trend.png",
        "caption": "Figure 3. Week 6 monthly ROI trend from the advertising dataset."
    },
    {
        "notebook": ROOT / "OMDS-MODB2-Week7-Elahi-Awais.ipynb",
        "match": "Highlighting the Final Three Months",
        "filename": "week7_highlighted_trend.png",
        "caption": "Figure 4. Week 7 preattentive highlighting of the final three months."
    },
    {
        "notebook": ROOT / "OMDS-MODB2-Week7-Elahi-Awais.ipynb",
        "match": "Brand Funnel Improvement After Campaign",
        "filename": "week7_storytelling_chart.png",
        "caption": "Figure 5. Week 7 storytelling chart of before/after funnel improvement."
    },
]


def extract_first_png_from_cell(cell):
    outputs = cell.get("outputs", [])
    for out in outputs:
        data = out.get("data", {}) if isinstance(out, dict) else {}
        img = data.get("image/png")
        if img:
            if isinstance(img, list):
                img = "".join(img)
            return base64.b64decode(img)
    return None


def find_cell_image(notebook_path: Path, source_match: str):
    nb = nbformat.read(notebook_path, as_version=4)
    for cell in nb.cells:
        if cell.get("cell_type") != "code":
            continue
        source = cell.get("source", "")
        if source_match in source:
            return extract_first_png_from_cell(cell)
    return None


saved_figures = []
for spec in FIG_SPECS:
    image_bytes = find_cell_image(spec["notebook"], spec["match"])
    if not image_bytes:
        continue
    out_path = FIG_DIR / spec["filename"]
    out_path.write_bytes(image_bytes)
    saved_figures.append((out_path, spec["caption"]))

if not REPORT_IN.exists():
    raise FileNotFoundError(f"Input report not found: {REPORT_IN}")

# Append figures to a copy of the final polished report.
doc = Document(REPORT_IN)
doc.add_page_break()

heading = doc.add_paragraph("Appendix: Selected Graphs from Weekly Assignments")
heading.runs[0].bold = True
heading.paragraph_format.first_line_indent = Inches(0)

intro = doc.add_paragraph(
    "The figures below are extracted from completed Week 5, Week 6, and Week 7 notebook outputs "
    "and included as visual evidence supporting the milestone analysis."
)
intro.paragraph_format.first_line_indent = Inches(0.5)

for fig_path, caption in saved_figures:
    doc.add_paragraph("")
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(fig_path), width=Inches(6.2))

    cap_p = doc.add_paragraph(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.first_line_indent = Inches(0)


doc.save(REPORT_OUT)
print(f"Created {REPORT_OUT}")
print(f"Included {len(saved_figures)} figures")
for fig_path, _ in saved_figures:
    print(f"- {fig_path.name}")
