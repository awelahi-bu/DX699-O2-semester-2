from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

output_path = "Milestone_Three_Rubric_Aligned_Report.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Milestone Three Report (Rubric-Aligned)")
title_run.bold = True
title_run.font.size = Pt(16)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("DX699 - AI for Leaders\nAwais Elahi\nProject: Bivariate and Time-Series Analysis of Social Media Advertising Performance")

doc.add_paragraph()

sections = [
    ("Project Statement and Description", [
        "This project evaluates campaign performance drivers in the Social Media Advertising dataset, with emphasis on how exploratory analysis can guide later supervised and unsupervised modeling.",
        "The business problem is to identify which campaign variables are most informative for ROI, detect data quality issues before modeling, and determine whether campaign outcomes are stable enough over time for reliable comparison.",
        "The expected impact is decision support for marketing strategy through evidence-based prioritization of channels, segments, and campaign goals."
    ]),
    ("Executive Summary", [
        "Using completed work from Weeks 5 to 7, this report applies preprocessing checks, univariate analysis, bivariate analysis, and storytelling-oriented visual design.",
        "Key findings: Engagement Score shows the strongest linear relationship with ROI; Clicks and Impressions appear partially redundant; monthly average ROI is relatively stable; and visual emphasis methods improve interpretability."
    ]),
    ("Exploratory Data Analysis: Preprocessing", [
        "Preprocessing included date parsing, currency-to-numeric conversion for Acquisition Cost, row/column integrity checks, and category balance checks.",
        "Observed dataset facts used in this report: 300,000 rows, 16 columns, and no blank fields in the checked records.",
        "Additional compatibility handling was performed in notebooks for pandas month-end frequencies using ME."
    ]),
    ("Exploratory Data Analysis: Univariate Analysis", [
        "Univariate inspection summarized center and range across key variables and reviewed distribution shapes in weekly visualizations.",
        "Observed ranges include ROI (0 to 8), Conversion Rate (0.01 to 0.15), Engagement Score (1 to 10), and broad scales for Clicks and Impressions.",
        "Implication: variable scaling and discretization patterns should be considered in later modeling."
    ]),
    ("Exploratory Data Analysis: Bivariate Analysis", [
        "Bivariate analysis included correlations, pair plots, scatter plots, and time-series trend checks.",
        "Correlation with ROI: Engagement Score (0.355), Clicks (0.188), Impressions (0.166), Conversion Rate (approx. 0), Acquisition Cost (approx. -0.002).",
        "Interpretation: engagement quality appears more informative than raw exposure volume, and some volume metrics are partially overlapping."
    ]),
    ("Data Conclusions: Expected and Unexpected Data", [
        "Expected: engagement-related variables would carry stronger performance signal; Clicks and Impressions would co-move.",
        "Unexpected: Conversion Rate and Acquisition Cost show very weak linear pairwise relationship with ROI at aggregate level.",
        "This highlights the limitation of relying on pairwise relationships alone."
    ]),
    ("Data Conclusions: Disqualified Data", [
        "No dataset was fully disqualified, but constrained-value variables were treated cautiously for direct inference.",
        "The primary dataset remained usable for EDA because of completeness and broad categorical coverage."
    ]),
    ("Data Conclusions: Supervised and Unsupervised Analyses", [
        "Suggested supervised approaches: regularized regression and tree-based regression for ROI, plus classification for high-vs-low ROI campaigns.",
        "Suggested unsupervised approaches: clustering campaign archetypes and dimensionality reduction to identify dominant variance patterns.",
        "Reasoning: weak pairwise signals and potential interactions suggest comparing linear baselines with nonlinear models."
    ]),
    ("Response to Dataset Creator Concerns", [
        "The dataset is oriented toward campaign performance evaluation and optimization.",
        "This EDA addresses those concerns by identifying practical ROI-linked signals, highlighting limits of naive metrics, and showing where segmentation can improve decision quality."
    ]),
    ("Weekly Graphs and Homework Completion", [
        "Weeks completed through Week 7 include full graph/homework completion for Week 5 and completed weekly graph, EDA, and storytelling tasks in Weeks 6 and 7.",
        "These notebook outputs provide evidence for the report claims."
    ]),
    ("Communication Quality", [
        "Charts and explanations are structured for a business audience using direct labeling, focused annotations, reduced clutter, and clear distinction between exploratory findings and causal claims."
    ]),
    ("Final Recommendations", [
        "1) Prioritize engagement features in subsequent modeling.",
        "2) Segment by channel, customer segment, and campaign goal.",
        "3) Compare interaction-aware models against linear baselines.",
        "4) Continue stakeholder-focused storytelling visuals for executive communication."
    ]),
]

for heading, paragraphs in sections:
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        p = doc.add_paragraph(paragraph)
        p.paragraph_format.line_spacing = 2

# Rubric coverage table
doc.add_heading("Rubric Coverage Map", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
header = table.rows[0].cells
header[0].text = "Milestone Three Criterion"
header[1].text = "Where Addressed in This Report"

rubric_rows = [
    ("Problem statement/description", "Project Statement and Description"),
    ("EDA preprocessing", "Exploratory Data Analysis: Preprocessing"),
    ("EDA univariate", "Exploratory Data Analysis: Univariate Analysis"),
    ("EDA bivariate", "Exploratory Data Analysis: Bivariate Analysis"),
    ("Expected/unexpected data", "Data Conclusions: Expected and Unexpected Data"),
    ("Disqualified data", "Data Conclusions: Disqualified Data"),
    ("Supervised/unsupervised analysis", "Data Conclusions: Supervised and Unsupervised Analyses"),
    ("Response to dataset creators", "Response to Dataset Creator Concerns"),
    ("Weekly graphs/homework", "Weekly Graphs and Homework Completion"),
    ("Communication", "Communication Quality")
]

for criterion, location in rubric_rows:
    row = table.add_row().cells
    row[0].text = criterion
    row[1].text = location

# References
doc.add_heading("References", level=1)
ref1 = doc.add_paragraph("Boston University. (2026). BU-omds-moduleB-MilestoneThree-guidelines-rubric.pdf.")
ref1.paragraph_format.line_spacing = 2
ref2 = doc.add_paragraph("Social_Media_Advertising.csv dataset used in weekly notebooks and milestone analysis.")
ref2.paragraph_format.line_spacing = 2

doc.save(output_path)
print(f"Created {output_path}")
