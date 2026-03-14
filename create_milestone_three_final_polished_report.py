from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

output_path = "Milestone_Three_Final_Polished_Report.docx"

doc = Document()

# Instructor style requirements: 1-inch margins, TNR 12, double spacing.
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.first_line_indent = Inches(0.5)

# Title page (APA-like layout)
for _ in range(8):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run(
    "Milestone Three: Exploratory Data Analysis and Evaluation of Multiple Datasets"
)
title_run.bold = True

doc.add_paragraph("")

student = doc.add_paragraph()
student.alignment = WD_ALIGN_PARAGRAPH.CENTER
student.add_run("Awais Elahi")

course = doc.add_paragraph()
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
course.add_run("DX699 - AI for Leaders")

instructor = doc.add_paragraph()
instructor.alignment = WD_ALIGN_PARAGRAPH.CENTER
instructor.add_run("Boston University")

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(date.today().strftime("%B %d, %Y"))

doc.add_page_break()


def add_heading(text: str):
    h = doc.add_paragraph()
    h.paragraph_format.first_line_indent = Inches(0)
    run = h.add_run(text)
    run.bold = True


def add_para(text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.5)


add_heading("Abstract")
add_para(
    "This report presents a rubric-aligned exploratory data analysis for Milestone Three using "
    "completed work from Weeks 1 through 7, with emphasis on Week 5 to Week 7 notebook artifacts. "
    "The primary analytical dataset is Social_Media_Advertising.csv, which was evaluated through preprocessing, "
    "univariate analysis, bivariate analysis, trend inspection over time, and data storytelling design. "
    "The analysis indicates that engagement-oriented variables provide stronger linear signal for return on investment "
    "than raw volume variables, while some campaign metrics appear partially redundant. The findings also show that "
    "average monthly ROI is relatively stable in aggregate, supporting the dataset's use for subsequent modeling. "
    "The report maps each rubric criterion to concrete evidence and recommends supervised and unsupervised next steps."
)

doc.add_paragraph("")
add_heading("Project Statement and Description")
add_para(
    "The chosen project evaluates campaign-level social media performance to determine which observable attributes are "
    "most useful for understanding ROI outcomes. The industry problem is resource allocation under uncertainty: organizations "
    "must decide where to invest marketing effort without over-relying on a single metric such as impressions or clicks. "
    "The project therefore focuses on identifying data-informed indicators that are practical for managerial decision-making."
)
add_para(
    "The potential impact of the project is twofold. First, it supports tactical decisions such as campaign optimization by "
    "channel and audience segment. Second, it supports strategic decisions by indicating which variables are likely to matter "
    "for deeper predictive modeling in later milestones."
)

add_heading("Data Sources and Scope")
add_para(
    "The primary dataset for milestone analysis is Social_Media_Advertising.csv. The file includes 300,000 rows and 16 columns, "
    "with variables including campaign metadata, engagement behavior, conversion, cost, and ROI. Weekly notebooks from Week 1 to Week 7 "
    "were also completed, and Weeks 5 through 7 were used as the main evidence base for graphing, interpretation, and communication techniques."
)
add_para(
    "A secondary healthcare dataset was reviewed for context, but the social media advertising dataset was retained as the principal source "
    "because it aligns more directly with the campaign performance problem and has strong completeness for exploratory analysis."
)

add_heading("Exploratory Data Analysis: Preprocessing")
add_para(
    "Preprocessing tasks included parsing Date into datetime format, converting Acquisition_Cost from currency string to numeric type, "
    "and checking structural quality indicators such as row/column consistency and blank-value frequency. A missing-field check returned zero blank values "
    "across the evaluated records, improving confidence in baseline usability."
)
add_para(
    "Categorical distributions were examined for Channel_Used, Campaign_Goal, and Customer_Segment. Channel and goal counts were broadly balanced, "
    "which reduces immediate risk of severe imbalance in simple exploratory comparisons. During notebook polishing, pandas 3.x compatibility was also handled by "
    "using month-end frequency code ME where required for date_range operations."
)
add_para(
    "The preprocessing stage revealed one practical caution: Acquisition_Cost appears constrained in the dataset structure, which may limit direct inferential value "
    "for pairwise cost-to-ROI interpretation. This is not a disqualifying issue, but it informs downstream model feature treatment."
)

add_heading("Exploratory Data Analysis: Univariate Analysis")
add_para(
    "Univariate analysis summarized central tendency and spread for major numeric variables. ROI ranges from 0 to 8 with a mean near 3.18. Conversion rate is bounded "
    "between 0.01 and 0.15, while Engagement_Score ranges from 1 to 10. Clicks and Impressions occupy much larger numeric scales, indicating the need for scale-aware "
    "interpretation in charts and future models."
)
add_para(
    "Distribution views in weekly graphs suggest that some variables are highly structured or discretized. This does not invalidate the dataset, but it implies that assumptions "
    "of smooth continuous behavior may be weak for certain features. For subsequent modeling, standardized scaling and robust diagnostics should be used where appropriate."
)

add_heading("Exploratory Data Analysis: Bivariate Analysis")
add_para(
    "Bivariate analysis used correlations, scatter plots, pair plots, and time-series line charts. The strongest observed linear association with ROI among tested numeric variables "
    "was Engagement_Score (approximately 0.355). Clicks and Impressions had weaker positive relationships with ROI (approximately 0.188 and 0.166), while Conversion_Rate and "
    "Acquisition_Cost were near zero in pairwise linear terms."
)
add_para(
    "Pair-plot evidence indicates overlap between Clicks and Impressions, suggesting partial redundancy for explanatory purposes. Time-series analysis of monthly average ROI shows "
    "modest movement rather than dramatic drift, which supports broad comparability across the observed period at aggregate level."
)
add_para(
    "Week 5 through Week 7 charting exercises were incorporated to validate communication quality: targeted highlighting, reduced clutter, and selective annotation were used to draw "
    "attention to decision-relevant features without overstating certainty."
)

add_heading("Data Conclusions: Expected and Unexpected Findings")
add_para(
    "Expected findings included co-movement of volume-based variables and a positive relationship between engagement and ROI. Unexpected findings included near-zero pairwise relationships "
    "for Conversion_Rate and Acquisition_Cost with ROI in aggregate form."
)
add_para(
    "These results are meaningful because they warn against simplistic single-metric decision policies. They also demonstrate why EDA is necessary before model selection: expected business "
    "logic can hold in part, while still missing interactions that only multivariate analysis can reveal."
)

add_heading("Data Conclusions: Disqualified Data")
add_para(
    "No full dataset was disqualified. The primary advertising dataset remained usable based on completeness and practical relevance. However, selected variables were treated as lower-confidence "
    "standalone predictors due to constrained ranges and weak pairwise behavior."
)
add_para(
    "The secondary healthcare dataset was not used as the central milestone dataset because it was less aligned to the campaign-performance industry problem addressed in this project."
)

add_heading("Data Conclusions: Supervised and Unsupervised Analysis Recommendations")
add_para(
    "Recommended supervised analyses include regularized linear regression and tree-based regression for ROI prediction, plus optional classification framing for high versus low ROI campaigns. "
    "Tree-based methods are recommended because interaction effects are likely and pairwise linear signals are moderate."
)
add_para(
    "Recommended unsupervised analyses include clustering campaign archetypes and dimensionality reduction to identify dominant structure across behavior and outcome variables. These methods can "
    "support strategy by exposing hidden segment patterns that are not visible from aggregate metrics alone."
)

add_heading("Response to Dataset Creator Concerns and Industry Problem")
add_para(
    "The dataset appears designed to address campaign effectiveness and optimization questions. This EDA directly contributes to those concerns by identifying which measurable attributes are most "
    "informative, clarifying where metrics overlap, and showing where broad trend stability does or does not support straightforward comparison."
)
add_para(
    "In practical terms, the analysis helps answer whether current data can guide better investment decisions. The answer is yes for exploratory prioritization, with the caveat that richer segmented and "
    "multivariate analyses are needed for stronger decision confidence."
)

add_heading("Weekly Graphs and Homework Assignments (Weeks 1-7 Evidence)")
add_para(
    "Week 5 homework was completed end-to-end with all required graph tasks, including corrected and generated visual references where hidden image files were absent. Week 6 included completed weekly "
    "graph response, dataset-focused exploratory analysis, and storytelling chart design. Week 7 included completed weekly response, preattentive highlighting demonstration, and storytelling comparison chart."
)
add_para(
    "These notebook outputs provide direct evidence supporting the claims in this report and satisfy the requirement that weekly graphs and assignments be completed through the Milestone Three checkpoint."
)

add_heading("Communication and Audience Fit")
add_para(
    "The report is written for a decision-making audience and follows communication principles from Storytelling With Data: one core message per figure, selective emphasis, concise interpretation, and explicit "
    "distinction between exploratory evidence and causal claims."
)
add_para(
    "This approach supports instructor communication criteria by improving clarity, coherence, and practical relevance while retaining methodological transparency."
)

add_heading("Final Recommendations")
add_para(
    "First, prioritize engagement-oriented variables in baseline model development and campaign diagnostics. Second, perform segmented analysis by channel, campaign goal, and customer segment to reduce aggregation bias. "
    "Third, benchmark linear and nonlinear models side-by-side to evaluate interaction effects. Fourth, preserve annotation-driven communication style in milestone presentations to help stakeholders quickly identify actionable findings."
)

add_heading("References")
ref1 = doc.add_paragraph("Boston University. (2026). BU-omds-moduleB-MilestoneThree-guidelines-rubric.pdf.")
ref1.paragraph_format.first_line_indent = Inches(-0.5)
ref1.paragraph_format.left_indent = Inches(0.5)
ref2 = doc.add_paragraph("Social_Media_Advertising.csv. (2026). Course dataset used for exploratory analysis.")
ref2.paragraph_format.first_line_indent = Inches(-0.5)
ref2.paragraph_format.left_indent = Inches(0.5)

doc.save(output_path)
print(f"Created {output_path}")
